import regex as re
from pypdf import PageObject, PdfReader
import os
import google.generativeai as genai
import json
import time
from docx import Document

# --- Configuration ---
# IMPORTANT: Replace with your actual API key or use a .env file for security.
os.environ["GOOGLE_API_KEY"] = "" 
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

PDF_INPUT_PATH = 'DepostionForPersisYu_LinkPDF.pdf'
DOCX_OUTPUT_PATH = 'toc.docx'
JSON_OUTPUT_PATH = 'topics.json'


# --- Helper Functions ---

def process_page(page: PageObject):
    """
    Takes a PDF page object, cleans the text, and returns the filtered text
    and the page number.
    """
    text = page.extract_text()
    if not text:
        return None, None

    # Use regex to find the page number first
    match = re.search(pattern=r"Page (\d+)", string=text)
    if not match:
        return None, None
    page_number = int(match.group(1))

    # Remove all timestamps (e.g., 10:32) from the text block
    text = re.sub(pattern=r'\b\d{2}:\d{2}\b', string=text, repl='')
    
    lines = text.split('\n')

    # Keep only lines that start with a line number 
    filtered_lines = [l for l in lines if re.search(r'^\d{1,2}\s', l)]
    new_text = '\n'.join(filtered_lines)

    return new_text, page_number

def find_line_number(text_block: str, starting_line: str):
    """
    Finds the line number for a given starting line of text within a block.
    """
    if not starting_line:
        return None
    for l in text_block.split('\n'):
        # Check if the cleaned line starts with the target text
        if l.strip().startswith(starting_line.strip()):
            line_num_str = l.split()[0]
            return int(line_num_str.strip())
    return None

def prompt_llm(last_topic: str, page_text: str) -> dict:
    """
    Sends a prompt to the Gemini model and returns the structured JSON response.
    """
    model = genai.GenerativeModel("gemini-2.0-flash")

    prompt = f"""
You are an expert paralegal analyzing a legal deposition. Your task is to identify when a new topic of discussion begins.

The last known topic was: "{last_topic}"

Here is the next page of the transcript:
---
{page_text}
---

If a new, distinct subject begins on this page, respond ONLY with the following JSON structure:
{{
  "boolTopicFound": true,
  "newtopicname": "A concise title for the new topic",
  "startinglineoftopic": "The exact first line of text where this new topic starts"
}}

If the discussion on this page is a continuation of the last known topic, respond ONLY with:
{{
  "boolTopicFound": false,
  "newtopicname": "",
  "startinglineoftopic": ""
}}

Do not add any explanations or introductory text. Return only the raw JSON object.
"""

    try:
        response = model.generate_content(prompt)
        llm_reply = response.text.strip()

        # Clean the response to ensure it's valid JSON
        json_start = llm_reply.find('{')
        json_end = llm_reply.rfind('}') + 1
        cleaned_json = llm_reply[json_start:json_end]

        return json.loads(cleaned_json)

    # In Python 3, standard exceptions are built-in and don't need to be imported.
    except Exception as e:
        print(f"An error occurred during the Gemini API call: {e}")
        # Return a default "not found" response on error
        return {'boolTopicFound': False}


# --- Main Execution ---

if __name__ == "__main__":
    lastTopic = 'Initial Appearances'
    topics = []

    try:
        reader = PdfReader(PDF_INPUT_PATH)
        number_of_pages = len(reader.pages)
        print(f"Found {number_of_pages} pages in '{PDF_INPUT_PATH}'.")

        for i in range(number_of_pages):
            page_object = reader.pages[i]
            
            # 1. Clean the page text
            newtext, page_number = process_page(page_object)

            # 2. If the page is valid, call the LLM
            if newtext and page_number:
                print(f"Processing Page: {page_number}...")
                output = prompt_llm(last_topic=lastTopic, page_text=newtext)
                
                # Pause for 1 second to respect the API's free tier rate limit
                time.sleep(1) 

                # 3. Process the LLM's output
                if output.get('boolTopicFound'):
                    linenum = find_line_number(newtext, output.get('startinglineoftopic', ''))

                    if linenum is not None:
                        new_topic_data = {
                            "topic": output['newtopicname'],
                            "page_start": page_number,
                            "line_start": linenum
                        }
                        topics.append(new_topic_data)
                        lastTopic = output['newtopicname']
                        print(f"  -> New topic found: {lastTopic}")

        print("\nProcessing complete.")

        # --- 4. Save results to a .docx file ---
        print(f"Saving Table of Contents to '{DOCX_OUTPUT_PATH}'...")
        document = Document()
        document.add_heading('Deposition Table of Contents', level=1)
        for item in topics:
            line = f"{item['topic']} ··· Page {item['page_start']} · Line {item['line_start']}"
            document.add_paragraph(line)
        document.save(DOCX_OUTPUT_PATH)
        print("Word document has been created.")

        # --- 5. Save raw data to a .json file for validation ---
        print(f"Saving raw data to '{JSON_OUTPUT_PATH}' for validation...")
        with open(JSON_OUTPUT_PATH, 'w') as f:
            json.dump(topics, f, indent=4)
        print("JSON file has been created.")

    except FileNotFoundError:
        print(f"Error: The file '{PDF_INPUT_PATH}' was not found.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
