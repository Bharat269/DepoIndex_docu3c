

import pypdf
import nltk
from nltk.tokenize import word_tokenize
from nltk.tag import pos_tag
from nltk.corpus import stopwords, wordnet
from nltk.stem import WordNetLemmatizer
import re
import argparse
import google.genai as genai
from google.genai import types
import json
from docx import Document
from typing import List, Dict, Tuple 
import pydantic

# NLTK Downloads
nltk.download('averaged_perceptron_tagger_eng')
nltk.download('stopwords')
nltk.download('punkt')
nltk.download('wordnet')
nltk.download('omw-1.4')

stopwords_set = set(stopwords.words('english'))

# Helper function for WordNetLemmatizer POS mapping 
def get_wordnet_pos(treebank_tag):
    if treebank_tag.startswith('J'):
        return wordnet.ADJ
    elif treebank_tag.startswith('V'):
        return wordnet.VERB
    elif treebank_tag.startswith('N'):
        return wordnet.NOUN
    elif treebank_tag.startswith('R'):
        return wordnet.ADV
    else:
        return wordnet.NOUN



def processText(path: str) -> Tuple[List[str], Dict[int, str]]:
    
    reader = pypdf.PdfReader(path)
    num_pages = reader.get_num_pages()
    
    # from a single string to a dictionary to hold processed pages
    processed_pages_dict = {}
    lemmatizer = WordNetLemmatizer()
    #Process each page and store it in the dictionary (Logic is mostly the same) ###
    for i in range(num_pages):
        end = False
        page = reader.get_page(i)
        text = page.extract_text()
        
        if text is None: # Handle cases where a page has no extractable text
            continue

        text = text.lower()
        if re.search(pattern='witness signature', string=text):
            end = True
        text = re.sub(pattern=r'\b\d{2}:\d{2}\b', string=text, repl='')
        text = re.sub(pattern=r"page (\d+)",repl='',string=text)

        tokens = word_tokenize(text)
                
        tagged_tokens_for_lemmatization = pos_tag(tokens)
        lemmatized_words = []
        for word, tag in tagged_tokens_for_lemmatization:
            lemmatized_words.append(lemmatizer.lemmatize(word, pos=get_wordnet_pos(tag)))
        
        tokens = lemmatized_words
        
        pos_tags = pos_tag(tokens)
        pos_to_remove = {'DT', 'IN', 'CC'}
        filtered_tokens_pos = [
            word for word, tag in pos_tags 
            if tag not in pos_to_remove and word not in stopwords_set and word not in {',',';'}
        ]
        
        # Reconstruct the text for the current page
        page_content_string = ' '.join(filtered_tokens_pos)
        page_content_string = '\n' + f"PAGE {i+1}" +'\n'+ page_content_string
        #  to store the processed page text in the dictionary
        # Page numbers are 1-based
        processed_pages_dict[i + 1] = page_content_string
        
        if end:
            break
    
    ### Create overlapping chunks from the processed pages ###
    chunked_list = []
    page_keys = sorted(processed_pages_dict.keys()) # Get page numbers 
    num_processed_pages = len(page_keys)
    
    chunk_size = 8
    step = 7 # This creates the 1-page overlap (chunk_size - overlap_size = 8 - 1 = 7)
    
    # Iterate through the page keys to create chunks
    for i in range(0, num_processed_pages, step):
        # Determine the page numbers for the current chunk
        chunk_end_index = i + chunk_size
        
        # Ensure we don't go past the last available page
        if chunk_end_index > num_processed_pages:
            # This logic will only create full chunks of 8.
            break

        # Get the keys (page numbers) for this chunk
        current_chunk_keys = page_keys[i:chunk_end_index]
        
        # Combine the text from the pages in this chunk
        chunk_content = ' '.join([processed_pages_dict[p] for p in current_chunk_keys])
        chunked_list.append(chunk_content)
    ## NEW SECTION END ###

    # return both the chunked list and the page dictionary
    return chunked_list, processed_pages_dict


def promptLLM(processed_Text:str):
    prompt = f"""You are DepoIndex, an expert paralegal that labels *new* topics in deposition transcripts. Always reply in JSON ONLY. 
    I have provided a lemmetized text, and removed stop words. The text is a complete deposition transcript. 
    Every page begins with PAGE (page number). Each line in the transcript starts with its line number (e.g., '1', '2', ..., '25'). Use this to find the starting point of each topic.
    Your job is to return a json in the structure following structure
    [  
    {{ "topic": "<string>",  
        "page_start": <int>,  
        "line_start": <int>  
    }},  
    ... 
    ]
    And i remind you, every page STARTS with the page number, every line starts with the line number. Re-use the **exact same topic label** (case-sensitive) if a new section appears similar in content or is a continuation of a previous topic. This maintains consistency in the output.
    These are a few examples.
    example 1 (
    input text = '''
        PAGE_1
        1 united state district court 2 southern district california 3 4 5 heather turrey oliver fiety ) jordan hernandez jeffrey ) 6 sazon individually ) behalf others ) 7 similarly situate ) ) 8 plaintiff ) ) 9 vs. ) case no. ) 3:20-cv-00697-dms ( ahg ) 10 vervent inc. fka first ) associate loan servicing ) 11 llc activate financial llc ) david johnson lawrence ) 12 chiavaro ) ) 13 defendant . ) ______________________________ ) 14 15 16 deposition persis yu 17 sacramento california 18 tuesday march 28 2023 19 20 21 22 report matthew sculatti csr . 13558 23 24 25 veritext legal solution 866 299-5127
        PAGE_2
        1 united state district court 2 southern district california 3 4 5 heather turrey oliver fiety ) jordan hernandez jeffrey ) 6 sazon individually ) behalf others ) 7 similarly situate ) ) 8 plaintiff ) ) 9 vs. ) case no. ) 3:20-cv-00697-dms ( ahg ) 10 vervent inc. fka first ) associate loan servicing ) 11 llc activate financial llc ) david johnson lawrence ) 12 chiavaro ) ) 13 defendant . ) ______________________________ ) 14 15 16 17 deposition persis yu take 18 matthew sculatti certified shorthand reporter 19 state california commence 1:15 p.m. 20 tuesday march 28 2023. deposition report 21 remotely veritext virtual technology . 22 23 24 25 veritext legal solution 866 299-5127
        PAGE_3
        1 appearance 2 3 4 plaintiff 5 blood hurst o'reardon llp 6 timothy g. blood attorney law 7 james m. davis attorney law 8 501 west broadway 9 suite 1490 10 san diego california 92101 11 tel ( 619 ) 338-1100 12 fax ( 619 ) 338-1101 13 tblood @ bholaw.com 14 jdavis @ bholaw.com 15 16 plaintiff 17 langer grogan diver pc 18 irv ackelsberg attorney law 19 1717 arch street 20 suite 4020 21 philadelphia pennsylvania 19103 22 tel ( 215 ) 320-5660 23 fax ( 215 ) 320-5703 24 iackelsberg @ langergrogan.com 25 veritext legal solution 866 299-5127'''
    output = [
        {{
        "topic": "Case title".
        "page_start": 1,  
        "line_start": 1 
        }},
        {{ 
        "topic": "Deposition particulars",  
        "page_start": 1,  
        "line_start": 5  
        }},
        {{ "topic": "Appearance for deposition",  
        "page_start": 3,  
        "line_start": 1  
        }}
        ]
    example 2 (
    input text = '''
        PAGE_22
        1 'll refer vervent 2 first associate time 2011 ? 3 aware data -- 'm sorry 4 say ? 5 q sure . 6 aware regulation 7 violate access group transfer data 8 peak loan client vervent 2011 ? 9 's scope ask 10 look case n't -- n't 11 access document transfer 12 could n't make opinion rule 13 violate . 14 q okay . ever job relate 15 criminal law ? 16 yes . intern 17 district attorney 's office 2002 . 18 q extent work experience 19 relating criminal law ? 20 -- tangentially 21 experience work survivor domestic violence . 22 q okay . experience 23 criminal law ? 24 volunteer teen court 25 attorney rochester . veritext legal solution 866 299-5127 
        PAGE_23
        1 q okay . ? 2 recall 3 instance . 4 serve jury . 5 q ever job decide 6 someone would charge crime ? 7 . 8 q familiar -- familiarity 9 rico statute ? 10 -- -- generally aware rico 11 statute . 12 q ever job help 13 determine rico enterprise exist ? 14 work -- consumer attorney -- 15 position consider work others 16 consider rico violation 17 occur . 18 q ever experience 19 context student loan ? 20 case would -- would 21 student loan relate case . 's primary 22 area expertise . 23 q many time ? 24 small handful time . less half -- 25 less half dozen . veritext legal solution 866 299-5127
        PAGE_24
        1 q last time ? 2 could n't recall . 3 q mean five ten year ago ? 4 ? 5 last five ten year sure . 6 q report attached hereto 7 first exhibit deposition make finding 8 rico enterprise ? 9 ask provide legal opinion 10 rico enterprise . 11 q make finding respect 12 issue ? 13 -- scope 14 report . 15 q go -- 'm go go back 16 report . 'm -- 's paragraph 8 . see 17 're ? paragraph 8 longer 18 paragraph number . conscious decision ? 19 conscious decision . 20 q okay . happen . 'm try give 21 hard time . wonder 's 22 significance first part prefatory 23 rest different somehow -- 24 ? okay . 25 . veritext legal solution 866 299-5127
        PAGE_25
        1 q okay . thank . 2 provide -- initial background 3 provide overview high education for-profit 4 college . familiarity concept 5 for-profit college america take ? 6 . 7 q -- 's law for-profit 8 school correct ? 9 correct . 10 q 's law make profit either 11 correct ? 12 correct . 13 q firsthand experience 14 quality education -- itt ? 15 mr. blood say firsthand ? 16 mr. purcell yes . 17 witness never attend itt . 18 mr. purcell 19 q -- -- part -- work 20 ever anything analyze -- -- 21 level -- education ? 22 -- area expertise student 23 lending area would 've examine 24 regard itt -- debt student 25 incur . -- -- individual veritext legal solution 866 299-5127
        PAGE_28
        1 level -- -- represent individual borrower -- 2 -- 've represent small number itt 3 student however expertise look 4 broadly itt systemic problem 5 identify regulator also senate 6 committee look systemically issue 7 raise student loan borrower 8 context understand way itt lead 9 student debt . 10 q would fair say view itt 11 base -- material 've review 12 many occasion benefit itt education 13 outweigh burden debt student take ? 14 'm sorry ask one time ? 15 q sure . 16 -- view many occasion itt 17 student receive less benefit education 18 would commensurate amount student 19 debt take ? 20 view -- let rephrase . 21 base analysis material 22 -- public domain report itt 23 report rely show 24 student -- student receive 25 benefit promise itt veritext legal solution 866 299-5127
        PAGE_29
        1 lead take debt first place . 2 q -- 's opinion itt 3 diploma worthless correct ? 4 ... 5 mr. blood vague . 6 witness well yeah . mean think 7 n't really understand mean say 8 -- worthless . mean data suggest 9 know folk attend school -- folk attend 10 for-profit college get diploma 11 itt typically earn less even 12 high school diploma . 13 think 's also issue 14 student complete degree . itt 15 for-profit school extraordinarily high 16 rate . 17 mr. purcell 18 q uh-huh . 19 many -- large -- systemically 20 large number folk -- even one say 21 -- degree value -- -- n't -- 22 know data earnings know 23 data retention . n't -- 'm position 24 say educational quality example . 25 -- say economic outcome veritext legal solution 866 299-5127
        PAGE_30
        1 student attend itt typic- -- -- 2 statistically bad many folk attend 3 itt . 4 mr. purcell 5 q okay . 'm try unpack little 6 bit -- 7 sure . 8 q -- -- 're talk 9 statistic . 10 -- statistic aware 11 many -- percentage itt student end 12 get degree oppose get 13 degree ? 14 'll -- -- 'm go refer report 15 believe date . believe -- 16 senate help committee look specifically itt 17 look retention rate itt . 18 mr. blood ms. yu -- senate 19 committee ? n't catch . 20 witness -- -- 'm sorry . 21 senate help committee health education 22 labor pension committee . 23 mr. blood thank . 24 witness report -- 25 release report 2012 base two-year veritext legal solution 866 299-5127
        PAGE_31
        1 investigation number for-profit institution . 2 itt one school investigate . 3 part investigation release data regard 4 retention rate itt . 'm happy 'll give 5 moment take look pull number 6 . 7 mr. purcell 8 q sure . 9 appear put exact figure 10 retention rate -- report . n't 11 document senate help report . -- 12 metric besides retention important look 13 -- judge quality degree 14 quality -- borrower experience -- 15 student rather experience valuable experience 16 look default rate institution itt 17 notably high default rate indicate -- 18 good indicator student get 19 value -- value education 20 take loan get . 21 q okay . couple sort big-picture question . 22 one -- 23 sure . 24 q -- one -- opinion 25 degree itt worthless fair ? veritext legal solution 866 299-5127'''
         
        output = [
        {{
        "topic": "Experience in criminal law".
        "page_start": 24,  
        "line_start": 14 
        }},
        {{ 
        "topic": "Knowledge about RICO statute",  
        "page_start": 25,  
        "line_start": 8  
        }},
        {{ "topic": "Experience in ITT institute(loans, student debt, education value)",  
        "page_start": 27,  
        "line_start": 13  
        }}
        ]
    Make sure to check for line number. It MUST be between 1 and 25. ITS GIVEN AT THE START OF EACH LINE.
    Page number will always be mentioned at start of each page. 
    
    Pay attention. This is crucial. I am giving you a chunk of pages. AVOID OVERLAP OF TOPICS. A TOPIC SHOULD NOT BE LISTED TWICE FOR TWO DIFFERENT PAGES.
    Now extract topics and return ONLY the JSON output.
    {processed_Text}
    """
    client = genai.Client(api_key="") #insert api key
    configuration = types.GenerateContentConfig(
        temperature=0.2,
        response_mime_type="application/json"    #to enforce JSON return
    )
    response = client.models.generate_content(
        model="gemini-2.5-flash-lite",
        contents=prompt,
        config=configuration
    )
    print("Response generated from LLM")
    
    # --- FIX ---
    # Convert the JSON string from the response's text into a Python list
    # This ensures a JSON-serializable object is returned, fixing the TypeError.
    try:
        return json.loads(response.text)
    except json.JSONDecodeError:
        print("   [!] Warning: Failed to decode JSON from the LLM response.")
        return [] # Return an empty list on failure to avoid crashing.

# STARTING FUNCTION CALLS
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Table of Contents for a deposition PDF")
    parser.add_argument("--file", type=str, required=True, help="Path to the input PDF file")
    parser.add_argument("--out", type=str, required=True, help="Base name for the output files (without extension)")

    args = parser.parse_args()

    print(f"1. Processing PDF: {args.file}")
    input_chunk_list, page_dictionary = processText(args.file)
    print(f"...Done. PDF processed into {len(input_chunk_list)} text chunks.")

    final_toc = []
    print("\n2. Generating topics from text chunks...")
    for text_chunk in input_chunk_list:
        if not text_chunk.strip():
            continue
        topics_from_chunk = promptLLM(text_chunk)
        if topics_from_chunk:
            final_toc.extend(topics_from_chunk)
    print(f"   ...Done. Aggregated a total of {len(final_toc)} topics.")

    # Check if any topics were generated before saving
    if not final_toc:
        print("\nNo topics were generated. Exiting without creating output files.")
        exit()

    print("\n3. Saving final Table of Contents in multiple formats...")

    # --- Save to JSON ---
    json_output_path = f"{args.out}.json"
    print(f"   -> Saving as JSON: {json_output_path}")
    with open(json_output_path, 'w') as json_file:
        json.dump(final_toc, json_file, indent=4)

    # --- Save to DOCX ---
    docx_output_path = f"{args.out}.docx"
    print(f"   -> Saving as DOCX: {docx_output_path}")
    doc = Document()
    doc.add_heading('Table of Contents', level=1)
    for item in final_toc:
        # Create a paragraph and add formatted runs (parts) to it
        p = doc.add_paragraph()
        p.add_run(f"{item['topic']}: ").bold = True
        p.add_run(f"Page {item['page_start']}, Line {item['line_start']}")
    doc.save(docx_output_path)

    # --- Save to Markdown ---
    md_output_path = f"{args.out}.md"
    print(f"   -> Saving as Markdown: {md_output_path}")
    with open(md_output_path, 'w') as md_file:
        md_file.write("# Table of Contents\n\n")
        for item in final_toc:
            line = f"- **{item['topic']}** (Page {item['page_start']}, Line {item['line_start']})\n"
            md_file.write(line)

    print("\nProcess complete! All files saved. ")